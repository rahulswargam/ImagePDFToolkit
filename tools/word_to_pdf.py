import os

from tools.image_resizer import get_output_folder
from tools.office_com import is_office_app_installed

_WD_FORMAT_PDF = 17


def _unique_pdf_path(output_folder, base_name, suffix):
    output_path = os.path.join(output_folder, f"{base_name}{suffix}.pdf")
    counter = 1
    while os.path.exists(output_path):
        output_path = os.path.join(output_folder, f"{base_name}{suffix}_{counter}.pdf")
        counter += 1
    return output_path


def _convert_via_com(input_path, output_path):
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(os.path.abspath(input_path), ReadOnly=True)
        try:
            document.SaveAs(output_path, FileFormat=_WD_FORMAT_PDF)
        finally:
            document.Close(False)
    finally:
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def _convert_via_fallback(input_path, output_path):
    if input_path.lower().endswith(".doc"):
        raise ValueError("Legacy .doc files need Microsoft Word installed to convert.")

    import docx
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table

    document = docx.Document(input_path)
    styles = getSampleStyleSheet()
    story = []

    for paragraph in document.paragraphs:
        style = styles["Heading1"] if paragraph.style.name.startswith("Heading") else styles["Normal"]
        story.append(Paragraph(paragraph.text or " ", style))
        story.append(Spacer(1, 6))

    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        story.append(Table(rows))
        story.append(Spacer(1, 12))

    if not story:
        raise ValueError("This document has no readable content.")

    SimpleDocTemplate(output_path, pagesize=LETTER).build(story)


def convert_word_to_pdf(input_path):
    """
    Convert a Word document into a PDF. Uses Microsoft Word (COM automation)
    when it's installed, for exact fidelity. Falls back to a simplified
    text/table renderer (no headers/footers, images, or complex formatting)
    when Word isn't available or the COM call fails.

    Returns:
        (output_path, method) where method is "com" or "fallback"
    """

    input_path = str(input_path)
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    output_path = _unique_pdf_path(get_output_folder(), base_name, "")

    if is_office_app_installed("word"):
        try:
            _convert_via_com(input_path, output_path)
            return output_path, "com"
        except Exception:
            pass

    _convert_via_fallback(input_path, output_path)
    return output_path, "fallback"
